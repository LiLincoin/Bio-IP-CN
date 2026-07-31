"""生化专利代理人助手

运行：streamlit run app.py
"""

from __future__ import annotations

import io
import re
from typing import Iterator

import streamlit as st
from docx import Document
from openai import OpenAI


st.set_page_config(
    page_title="生化专利代理人助手",
    page_icon="🧬",
    layout="wide",
)


SYSTEM_PROMPT = """你是一位资深的中国生化领域专利代理人，负责协助撰写中国发明专利申请文件草稿。

请严格遵守以下规则：
1. 仅依据技术交底书中明确提供的事实、数据和实验条件进行撰写；不得捏造实验数据、文献、序列、化合物结构、菌种保藏信息或技术效果。
2. 对交底书未提供但专利撰写需要的信息，使用【待补充：具体信息】标注，并在文末列出待补充事项。
3. 权利要求书应具有清楚、简要、以说明书为依据的层次结构；合理区分产品、组合物、用途和方法权利要求，避免不必要的功能性或结果性限定。
4. 对化合物、核酸、蛋白质、抗体、微生物、细胞、检测试剂或生物制品，保留并准确使用交底书中的名称、编号、序列、保藏号、浓度、宿主和实验条件；没有依据时不得自行补写。
5. 说明书应包含：发明名称、摘要、权利要求书、技术领域、背景技术、原理说明、发明内容、附图说明（如有）、具体实施方式和待补充事项。
6. “原理说明”必须单独成节：结合交底书明确给出的生物化学机制、反应过程、信号通路、结构-功能关系或检测原理，解释技术方案为什么能够解决技术问题；不得把推测写成事实，缺少依据处标注【待补充】。
7. 背景技术应客观表述，不能虚构具体文献或现有技术结论。摘要控制在约300字以内。
8. 实施例要清楚记载材料、步骤、对照、检测方法、统计方法和结果；对于缺失内容明确标注。
9. 输出中文 Markdown，不要输出与专利文件无关的闲聊。先给出“事实依据与风险提示”，再给出完整专利申请文件草稿。
10. 明确提示：该结果是撰写辅助草稿，必须由专利代理师和技术人员审核，不能替代法律意见或实验验证。
"""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """提取 DOCX 中的段落和表格文本。"""
    document = Document(io.BytesIO(file_bytes))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        blocks.append(f"[表格 {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", "；") for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))

    return "\n".join(blocks)


def extract_text_from_file(file_name: str, file_bytes: bytes) -> str:
    """按扩展名读取 DOCX、TXT、Markdown 等技术交底材料。"""
    suffix = file_name.lower().rsplit(".", 1)[-1] if "." in file_name else ""
    if suffix == "docx":
        return extract_text_from_docx(file_bytes)
    if suffix in {"txt", "md", "markdown", "rst", "csv", "json"}:
        for encoding in ("utf-8-sig", "gb18030"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("文本文件不是 UTF-8 或 GB18030 编码。")
    raise ValueError(f"暂不支持 .{suffix or '未知'} 文件格式。")


def markdown_to_docx(markdown_text: str) -> bytes:
    """将模型生成的 Markdown 草稿转换为可下载的 DOCX 文件。"""
    document = Document()
    document.core_properties.title = "生化专利申请撰写辅助草稿"
    document.add_heading("生化专利申请撰写辅助草稿", level=0)
    document.add_paragraph("说明：本文档由 AI 生成，仅供专利代理师和技术人员审核，不构成法律意见。")

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^#{1,6}\s+(.+)$", line)
        if heading:
            level = min(len(raw_line) - len(raw_line.lstrip("#")), 9)
            document.add_heading(heading.group(1).strip(), level=level)
        elif re.match(r"^[-*+]\s+", line):
            document.add_paragraph(re.sub(r"^[-*+]\s+", "", line), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", line):
            document.add_paragraph(re.sub(r"^\d+[.)]\s+", "", line), style="List Number")
        else:
            document.add_paragraph(line)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def stream_patent_draft(
    api_key: str,
    base_url: str,
    model: str,
    input_text: str,
    temperature: float,
) -> Iterator[str]:
    """调用 OpenAI 兼容接口并逐段返回模型输出。"""
    client_kwargs = {"api_key": api_key.strip()}
    if base_url.strip():
        client_kwargs["base_url"] = base_url.strip()
    client = OpenAI(**client_kwargs)

    user_prompt = f"""以下是申请人提供的技术交底书：

--- 技术交底书开始 ---
{input_text}
--- 技术交底书结束 ---

请根据上述材料生成专利撰写辅助草稿。所有没有事实依据的内容都必须标记为【待补充】或列入待补充事项。"""

    response = client.chat.completions.create(
        model=model.strip(),
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        temperature=temperature,
        stream=True,
    )

    for chunk in response:
        if not chunk.choices:
            continue
        content = chunk.choices[0].delta.content
        if content:
            yield content


st.title("🧬 生化专利代理人助手")
st.caption("上传 DOCX、TXT、MD 等技术交底材料，使用您自己的 OpenAI 兼容 API 生成中国发明专利撰写辅助草稿。")

with st.sidebar:
    st.header("⚙️ API 配置")
    api_key = st.text_input(
        "API Key",
        type="password",
        help="仅用于当前会话中的 API 调用；请勿将密钥写入代码或提交到版本库。",
    )
    base_url = st.text_input(
        "API Base URL（可选）",
        value="https://api.openai.com/v1",
        help="也支持 DeepSeek、通义千问等提供 OpenAI 兼容接口的服务。",
    )
    model_name = st.text_input("模型名称", value="gpt-4o-mini")
    temperature = st.slider("生成随机性", 0.0, 1.0, 0.2, 0.1)
    st.divider()
    st.warning(
        "生成内容仅为辅助草稿。请由专利代理师和技术人员核验权利要求、" 
        "充分公开、序列/结构、实验数据及法律合规性。"
    )

uploaded_file = st.file_uploader(
    "📂 上传技术交底书或技术材料",
    type=["docx", "txt", "md", "markdown", "rst", "csv", "json"],
    help="支持 DOCX、TXT、Markdown、RST、CSV 和 JSON；文本文件支持 UTF-8 或 GB18030 编码。",
)

if uploaded_file is not None:
    try:
        document_text = extract_text_from_file(uploaded_file.name, uploaded_file.getvalue())
    except Exception as exc:
        st.error(f"文件解析失败：{exc}")
        st.stop()

    if not document_text.strip():
        st.error("未提取到正文或表格文本，请确认 DOCX 文件包含可读取内容。")
        st.stop()

    st.success(f"已读取：{uploaded_file.name}（约 {len(document_text):,} 个字符）")
    with st.expander("查看提取内容", expanded=False):
        st.text_area("技术交底书文本", document_text, height=320, label_visibility="collapsed")

    if st.button("🚀 开始撰写专利", type="primary", use_container_width=True):
        if not api_key.strip():
            st.warning("请先在左侧输入 API Key。")
        elif not model_name.strip():
            st.warning("请填写模型名称。")
        else:
            output_placeholder = st.empty()
            full_draft = ""
            try:
                with st.spinner("正在生成专利撰写辅助草稿，请稍候……"):
                    for token in stream_patent_draft(
                        api_key, base_url, model_name, document_text, temperature
                    ):
                        full_draft += token
                        output_placeholder.markdown(full_draft)

                if full_draft:
                    st.success("撰写完成。")
                    col_md, col_docx = st.columns(2)
                    with col_md:
                        st.download_button(
                            "📥 下载 Markdown 草稿",
                            data=full_draft.encode("utf-8"),
                            file_name="生化专利申请撰写辅助草稿.md",
                            mime="text/markdown",
                            use_container_width=True,
                        )
                    with col_docx:
                        st.download_button(
                            "📄 下载 DOCX 草稿（含原理说明）",
                            data=markdown_to_docx(full_draft),
                            file_name="生化专利申请撰写辅助草稿.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                else:
                    st.warning("模型未返回文本内容。")
            except Exception as exc:
                st.error(f"API 调用失败：{exc}")
else:
    st.info("请上传 DOCX、TXT、MD 等技术材料后开始。")
